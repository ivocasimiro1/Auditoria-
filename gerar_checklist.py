from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)

# ── Styles helpers
def add_heading(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x06, 0x0D, 0x1B)  # navy
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0xC9, 0xA0, 0x2A)  # gold
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(2)
    return p

def add_body(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    return p

def add_checkbox(text, indent_level=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.5 + indent_level * 0.5)
    run = p.add_run("☐  " + text)
    run.font.size = Pt(10.5)
    return p

def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.italic = True
    return p

# ═══════════════════════════════════════════════
# TITLE BLOCK
# ═══════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run("🏄  Checklist de Dados — Nova Escola de Surf")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x06, 0x0D, 0x1B)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run("Preenche e envia — o site fica pronto em minutos.")
r2.font.size = Pt(11)
r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
p_sub.paragraph_format.space_after = Pt(6)

doc.add_paragraph()  # spacer

# ═══════════════════════════════════════════════
# 1. IDENTIDADE
# ═══════════════════════════════════════════════
add_heading("1. Identidade da Escola", 2)
add_checkbox("Nome da escola  →  _________________________________")
add_checkbox("Slogan / tagline  →  _________________________________")
add_checkbox("Ano de fundação  →  _______")
add_checkbox("Descrição curta (2–3 frases sobre a escola e a sua filosofia)")
add_checkbox("Logótipo  (ficheiro PNG ou SVG, fundo transparente)")
add_checkbox("Foto de capa / hero  (landscape, mín. 1920 × 1080 px)")

# ═══════════════════════════════════════════════
# 2. LOCALIZAÇÃO & CONTACTO
# ═══════════════════════════════════════════════
add_heading("2. Localização & Contacto", 2)
add_checkbox("Morada completa  →  _________________________________")
add_checkbox("Número de WhatsApp  (com prefixo país, ex: 351912345678)  →  ___________")
add_checkbox("Email de contacto  →  _________________________________")
add_checkbox("Link Google Maps  ou  coordenadas GPS  (lat, lon)")
add_checkbox("Domínio próprio  (ex: arrifanasurfacademy.com)  →  ___________")
add_note("Nota: o domínio pode ficar a apontar para Cloudflare Pages — gratuito, sem custo mensal.")

# ═══════════════════════════════════════════════
# 3. PRAIAS
# ═══════════════════════════════════════════════
add_heading("3. Praias onde a escola opera", 2)
add_note("Preencher para cada praia (mínimo 1, máximo 6).")
for i in range(1, 5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(f"Praia {i}")
    r.bold = True
    r.font.size = Pt(10.5)
    add_checkbox(f"Nome  →  _________________________________", 1)
    add_checkbox(f"Região / município  →  _________________________________", 1)
    add_checkbox(f"Coordenadas GPS  lat _______ lon _______", 1)
    add_checkbox(f"Descrição curta (nível, tipo de fundo, vento ideal)", 1)

# ═══════════════════════════════════════════════
# 4. MODALIDADES E PREÇOS
# ═══════════════════════════════════════════════
add_heading("4. Modalidades e Preços", 2)
add_note("Preencher para cada tipo de aula oferecida.")
for i in range(1, 5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(f"Aula {i}")
    r.bold = True
    r.font.size = Pt(10.5)
    add_checkbox(f"Nome  (ex: Iniciante, Privada, Bodyboard)  →  _________________", 1)
    add_checkbox(f"Preço  →  € _______   por:  ☐ pessoa   ☐ sessão", 1)
    add_checkbox(f"Máx. de alunos  →  _______    Duração  →  _______ h", 1)
    add_checkbox(f"O que está incluído  (fato, prancha, seguro, vídeo…)", 1)

# ═══════════════════════════════════════════════
# 5. HORÁRIOS
# ═══════════════════════════════════════════════
add_heading("5. Horários e Época", 2)
add_checkbox("Horários disponíveis  (ex: 9h, 11h, 14h, 16h)  →  _________________________________")
add_checkbox("Dias da semana operacionais  (ex: Segunda a Domingo)  →  _________________")
add_checkbox("Meses de época  (ex: Abril – Outubro)  →  _________________")

# ═══════════════════════════════════════════════
# 6. EQUIPA / INSTRUTORES
# ═══════════════════════════════════════════════
add_heading("6. Equipa / Instrutores", 2)
add_note("Até 3 instrutores destacados na página. Podes ter mais na equipa.")
for i in range(1, 4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(f"Instrutor {i}")
    r.bold = True
    r.font.size = Pt(10.5)
    add_checkbox(f"Nome completo  →  _________________________________", 1)
    add_checkbox(f"Foto  (quadrada ou retrato, mín. 400 × 400 px)", 1)
    add_checkbox(f"Bio curta (1 linha)  →  _________________________________", 1)
    add_checkbox(f"Tags de especialidade  (ex: Iniciante, Avançado, Surf feminino)", 1)

# ═══════════════════════════════════════════════
# 7. ESTATÍSTICAS / CREDENCIAIS
# ═══════════════════════════════════════════════
add_heading("7. Estatísticas & Credenciais", 2)
add_checkbox("Nº de alunos por ano  (ex: 800+)  →  _______")
add_checkbox("Anos de experiência  →  _______")
add_checkbox("% de satisfação  (ex: 98%)  →  _______")
add_checkbox("Certificação principal  (ex: ISA, FPS)  →  _______")

# ═══════════════════════════════════════════════
# 8. AVALIAÇÕES
# ═══════════════════════════════════════════════
add_heading("8. Avaliações de Clientes", 2)
add_note("3 a 5 reviews reais. Quanto mais autênticas melhor — nome, país e texto.")
for i in range(1, 5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(f"Review {i}")
    r.bold = True
    r.font.size = Pt(10.5)
    add_checkbox(f"Nome do cliente  →  _________________________________", 1)
    add_checkbox(f"País / bandeira  (ex: 🇩🇪 Alemanha)  →  _________________", 1)
    add_checkbox(f"Texto da avaliação  (PT ou EN, 1–3 frases)", 1)

# ═══════════════════════════════════════════════
# 9. GALERIA
# ═══════════════════════════════════════════════
add_heading("9. Galeria de Fotos", 2)
add_checkbox("6 a 12 fotos da escola, praias e aulas  (formato 4:3 ou 16:9, mín. 800 px largura)")
add_note("Fotos com pessoas em ação funcionam muito melhor do que paisagens vazias.")

# ═══════════════════════════════════════════════
# 10. CONFIGURAÇÃO TÉCNICA
# ═══════════════════════════════════════════════
add_heading("10. Configuração Técnica", 2)
add_checkbox("Identificador único da escola  (ex: arrifana-surf-academy, nazare-surf)  →  ___________")
add_note("Será usado internamente para separar os dados de cada escola.")
add_checkbox("Domínio próprio  →  _________________________________")
add_note("Aponta o domínio para Cloudflare Pages — é gratuito, rápido e sem custo mensal.\nO mesmo projeto Supabase é partilhado por todas as escolas (separação por school_id).")

# ═══════════════════════════════════════════════
# RODAPÉ
# ═══════════════════════════════════════════════
doc.add_paragraph()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = p_footer.add_run("Envia este ficheiro preenchido e o site fica configurado em minutos.  🤙")
r_f.font.size = Pt(10)
r_f.font.color.rgb = RGBColor(0xC9, 0xA0, 0x2A)
r_f.bold = True

out_path = "/home/user/Auditoria-/Checklist_Nova_Escola_Surf.docx"
doc.save(out_path)
print("Saved:", out_path)
